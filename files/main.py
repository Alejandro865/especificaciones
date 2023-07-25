import os
import shutil
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
from docx.shared import Inches

################## CONFIGURACION DE USUARIO ####################
# Documentation:
# https://python-docx.readthedocs.io/en/latest/index.html
#https://pybonacci.org/2020/06/30/escribiendo-ficheros-docx-de-word-con-python-capitulo-vii-estilos/
# Output Path 
OUTPUT_PATH = '.\Outputs'

# Excel File Path
EXCEL_PATH = '.\Input\data.xlsx'

# Word File Path -- Template --
PLANTILLA_PATH ='.\Input\Template\plantilla.docx'

# Imeges Path
IMAGE_PATH = './Input/Images'

################## CONFIGURACION DE USUARIO ####################

# Routine to delete and create output dir.
def EliminarCrearCarpetas(path):
  # Check if the folder exists and deleted
  if(os.path.exists(path)):
    shutil.rmtree(path)
  # Create output dir only after delete it
    os.mkdir(path)


# Rutine to read excel data
def LeerDatos(path, worksheet):
  excel_df = pd.read_excel(path, worksheet)

  return excel_df


# Processing the template
l_tpl = PLANTILLA_PATH
docx_tpl = Document(PLANTILLA_PATH)


# Rutine to create a chapter 

def CrearCapitulo(df_datos,nombre_cap):
   
    docx_tpl.add_heading("Especificacion de"+" "+nombre_cap, level=1)
    for r_idx, r_val in df_datos.iterrows():
                            
        if (r_val['Incluir'] == 'SI' and r_val['Capitulo'] == nombre_cap):
            docx_tpl.add_paragraph('')
            docx_tpl.add_paragraph(r_val['Nombre'], style="Subtitle")
            docx_tpl.add_paragraph(r_val['Descripcion'])
            docx_tpl.add_paragraph('Medida', style="BODY_TEXT")
            docx_tpl.add_paragraph(r_val['Medida'], style='Cuerpo')
            docx_tpl.add_paragraph('Imagen', style="BODY_TEXT")
            img_path = IMAGE_PATH +'//'+r_val['Imagen']
            docx_tpl.add_picture(img_path, width=Inches(1.7))
            
            docx_tpl.add_page_break()
        elif (r_val['Incluir'] == 'NO'):
            pass

# Rutiona para crear Documento

def CrearWord():

  # Se adiciona un parrafo

   # docx_tpl.render(context)
    df_datos = LeerDatos(EXCEL_PATH, 'DATA')
   #  Using the function 'CrearCapitulo' you can create the chapter tha need 
   # in the specifications documents

    CrearCapitulo(df_datos,'Tuberia')
    CrearCapitulo(df_datos,'Cable')
    # CrearCapitulo(df_datos,'capitulo')

     #Guardar el documento
    nombre_doc = 'Especificaciones.docx'
    docx_tpl.save(OUTPUT_PATH + '\\' + nombre_doc)
    #===========================================================================
   

# Rutina principal
def run():
 # Eliminar y volver a crear la carpeta outputs
  EliminarCrearCarpetas(OUTPUT_PATH)

 # Crear word
  CrearWord()

if __name__ == '__main__':
  run()