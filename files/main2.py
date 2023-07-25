import os
import shutil
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
from docx.shared import Inches

################## CONFIGURACION DE USUARIO ####################
# https://python-docx.readthedocs.io/en/latest/index.html
# Ruta de salida 
OUTPUT_PATH = '.\Outputs'

# Ruta fichero Excel
EXCEL_PATH = '.\Input\data.xlsx'

# Ruta plantilla fichero word
PLANTILLA_PATH ='.\Input\Template\especificaciones_template_1.docx'

# Ruta de Imagenes
IMAGE_PATH = './Input/Images'

################## CONFIGURACION DE USUARIO ####################

# Rutina para eliminar y crear carpeta
def EliminarCrearCarpetas(path):
  # Verificar si la carpeta existe y eliminarla
  if(os.path.exists(path)):
    shutil.rmtree(path)
  # Crear Carpetya de salida
    os.mkdir(path)


# Rutina para leer datos de excel
def LeerDatos(path, worksheet):
  excel_df = pd.read_excel(path, worksheet)

  return excel_df

# Rutina para crear el fichero Word para cada material
def CrearWord(df_datos):
  for r_idx, r_val in df_datos.iterrows():
    # Carga plantilla
    if (r_val['Incluir'] == 'SI'):
      l_tpl = PLANTILLA_PATH
    elif (r_val['Incluir'] == 'NO'):
      pass
    
    # Procesando la plantilla
    docx_tpl = DocxTemplate(l_tpl)

    # Anador imagen grafica circular y de barras
    img_path = IMAGE_PATH +'\\'+r_val['Imagen']
    img = InlineImage(docx_tpl, img_path, height=Mm(15))
   
  # Crear contexto
    context = {
      'name': r_val['Nombre'],
      'description': r_val['Descripcion'],
      'pay': r_val['Medida'],
      'picture': img,
    }

  #Renderizar usando contexto
  docx_tpl.render(context)
  # Se adiciona un parrafo

  prueba = "a"
  docx_tpl.add_heading('Heading, level 2', level=2)
  docx_tpl.add_paragraph('')

  def CrearCapitulo(nombre_cap,df):  
    for r_idx, r_val in df.iterrows():
      if (r_val['Incluir'] == 'SI' and r_val['Capitulo'] == nombre_cap):
        docx_tpl.add_paragraph(r_val['Imagen'])
      elif (r_val['Incluir'] == 'NO'):
        pass
 


  # Adicionar una tabla

  records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
  )

  table = docx_tpl.add_table(rows=1, cols=3)
  hdr_cells = table.rows[0].cells
  hdr_cells[0].text = 'Qty'
  hdr_cells[1].text = 'Id'
  hdr_cells[2].text = 'Desc'
  for qty, id, desc in records:
      row_cells = table.add_row().cells
      row_cells[0].text = str(qty)
      row_cells[1].text = id
      row_cells[2].text = desc

  #Guardar el documento
  if (pd.notna(r_val['Nombre'])):
    nombre_doc = 'Especificaciones2.docx'
    docx_tpl.save(OUTPUT_PATH + '\\' + nombre_doc)



# Rutina principal
def run():
 # Eliminar y volver a crear la carpeta outputs
  EliminarCrearCarpetas(OUTPUT_PATH)

 # Leer datos del fichero excel
  df_datos = LeerDatos(EXCEL_PATH, 'DATA') 
 
 # Crear word
  CrearWord(df_datos)




if __name__ == '__main__':
  run()
