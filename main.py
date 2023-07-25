import os
import shutil
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
from docx.shared import Inches

################## SET-UP ####################
# Documentation:
# https://python-docx.readthedocs.io/en/latest/index.html
# https://pybonacci.org/2020/06/30/escribiendo-ficheros-docx-de-word-con-python-capitulo-vii-estilos/

# Output Path 
OUTPUT_PATH = '.\Outputs'

# Excel File Path
EXCEL_PATH = '.\Input\data.xlsx'

# Word File Path -- Template --
PLANTILLA_PATH ='.\Input\Template\plantilla.docx'

# Imeges Path
IMAGE_PATH = './Input/Images'

################## SET-UP ####################

# Routine to delete and create output dir.
def EliminarCrearCarpetas(path):
  # Check if the folder exists and deleted
  if(os.path.exists(path)):
    shutil.rmtree(path)
  # Create output dir only after delete it
    os.mkdir(path)


# Rutine to read excel data
def LeerDatos(path, worksheet):
  excel_df = pd.read_excel(path, worksheet)
  return excel_df

# Rutine to create the Specifications File in word
def CrearWord(df_datos):
  for r_idx, r_val in df_datos.iterrows():
           
    # Processing Template
    docx_tpl = DocxTemplate(PLANTILLA_PATH)

    # Rutine for created image to use in the context
    img_path = IMAGE_PATH +'\\'+r_val['Imagen']
    img = InlineImage(docx_tpl, img_path, height=Mm(15))
   
  # Create the context -- Values to remplace in the template
    context = {
      'Nombre_del_proyecto': 'Proyecto Ejemplo',
      'description': r_val['Descripcion'],
      'pay': r_val['Medida'],
      'picture': img,
    }

  #Render context
  docx_tpl.render(context)
 
#===================================================================
# Function to create the charp for diferente kinds of material 
  def CrearCapitulo(df_datos,nombre_cap):
   
    docx_tpl.add_heading("Especificación de"+" "+nombre_cap, level=1)
    for r_idx, r_val in df_datos.iterrows():
                            
        if (r_val['Incluir'] == 'SI' and r_val['Capitulo'] == nombre_cap):
            docx_tpl.add_paragraph('')
            docx_tpl.add_paragraph(r_val['Nombre'], style="Subtitle")
            docx_tpl.add_paragraph(r_val['Descripcion'])
            docx_tpl.add_paragraph('Medida y forma de pago', style="BODY_TEXT")
            docx_tpl.add_paragraph(r_val['Medida'], style='Cuerpo')
            docx_tpl.add_paragraph('Imagen', style="BODY_TEXT")
            img_path = IMAGE_PATH +'//'+r_val['Imagen']
            docx_tpl.add_picture(img_path, width=Inches(1.7))
            
            docx_tpl.add_page_break()
        elif (r_val['Incluir'] == 'NO'):
            pass
#===================================================================
# You must select the charpers to create and use de funtion CrearCapitulo
  CrearCapitulo(df_datos,'Transformador')
  CrearCapitulo(df_datos,'Ducteria')
  CrearCapitulo(df_datos,'Cable')
# CrearCapitulo(df_datos,'capitulo')

# Guardar el documento
  if (pd.notna(r_val['Nombre'])):
    nombre_doc = 'Especificaciones.docx'
    docx_tpl.save(OUTPUT_PATH + '\\' + nombre_doc)

# Main Rutine 
def run():
 # Delete and Create the folder
  EliminarCrearCarpetas(OUTPUT_PATH)

 # Read Excel File
  df_datos = LeerDatos(EXCEL_PATH, 'DATA') 
 
 # Create word
  CrearWord(df_datos)

if __name__ == '__main__':
  run()
